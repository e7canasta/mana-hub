#!/usr/bin/env python3
"""Arma .docx de seguimiento a partir del JSON de insights.

El hub/panel pide un informe por API; no sabe si se generó antes o ahora.
Este script consume el mismo contrato JSON. El envío por email viene después.

  python scripts/reports/generate.py --resident jose --days 30
  python scripts/reports/generate.py --facility --days 14

Requiere insights en :8081 y: pip install -r scripts/reports/requirements.txt
"""

from __future__ import annotations

import argparse
import json
import os
import sys
import urllib.error
import urllib.request
from datetime import datetime, timezone
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("Falta python-docx. pip install -r scripts/reports/requirements.txt", file=sys.stderr)
    sys.exit(1)

ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "generated"
DEFAULT_INSIGHTS = os.environ.get("INSIGHTS_URL", "http://localhost:8081")


def get_json(base: str, path: str) -> dict:
    url = f"{base.rstrip('/')}{path}"
    req = urllib.request.Request(url, headers={"Accept": "application/json"})
    try:
        with urllib.request.urlopen(req, timeout=60) as resp:
            return json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="replace")
        raise SystemExit(f"GET {path} -> {exc.code}: {detail[:400]}") from exc
    except urllib.error.URLError as exc:
        raise SystemExit(f"No pude hablar con insights en {url}: {exc.reason}") from exc


def _heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def _para(doc: Document, text: str, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.italic = italic


def _footer(doc: Document) -> None:
    _para(
        doc,
        "Este informe lo arma insights a partir del cubo diario de la residencia. "
        "No es un diagnóstico. Al aplicar un cambio de alarma, el hallazgo queda como motivo "
        "junto con los episodios que lo originaron.",
        italic=True,
    )


def resident_doc(payload: dict) -> Document:
    doc = Document()
    name = payload.get("residentName") or payload.get("residentId") or "Residente"
    _heading(doc, f"Informe de seguimiento — {name}", 0)
    _para(
        doc,
        f"Ventana {payload.get('from')} → {payload.get('to')} · "
        f"generado {payload.get('generatedAt', '')}",
    )

    _heading(doc, "Lo que avisa hoy", 1)
    lines = payload.get("policyToday") or []
    if not lines:
        _para(doc, "Sin perfil de alarma cargado para este residente.")
    for line in lines:
        doc.add_paragraph(line, style="List Bullet")

    _heading(doc, "Sueño", 1)
    if payload.get("narrative"):
        _para(doc, payload["narrative"])
    cards = payload.get("sleepCards") or []
    if cards:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Indicador"
        hdr[1].text = "Valor"
        for card in cards:
            row = table.add_row().cells
            row[0].text = card.get("label") or card.get("code") or ""
            detail = card.get("detail")
            value = card.get("value") or ""
            row[1].text = f"{value} ({detail})" if detail else value

    _heading(doc, "Hallazgos", 1)
    findings = payload.get("findings") or []
    if not findings:
        _para(doc, "Sin hallazgos en esta ventana.")
    for finding in findings:
        title = finding.get("headline") or finding.get("code") or "Hallazgo"
        if finding.get("awaitingDecision"):
            title = f"Hay una decisión esperándote: {title}"
        _heading(doc, title, 2)
        if finding.get("body"):
            _para(doc, finding["body"])
        proposal = finding.get("proposal")
        if proposal:
            _para(doc, f"Propuesta: {proposal.get('text', '')}")
            _para(
                doc,
                f"{proposal.get('applyLabel', 'Aplicar el cambio')} / "
                f"{proposal.get('dismissLabel', 'No hacerlo')}",
            )

    _heading(doc, "Episodios", 1)
    episodes = payload.get("episodes") or []
    if not episodes:
        _para(doc, "Sin episodios en la ventana.")
    else:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Cuándo"
        hdr[1].text = "Tipo"
        hdr[2].text = "Severidad"
        hdr[3].text = "Cierre"
        for ep in episodes:
            row = table.add_row().cells
            row[0].text = str(ep.get("occurredAt") or "")
            row[1].text = str(ep.get("kind") or "")
            row[2].text = str(ep.get("severity") or "")
            if ep.get("selfRecovery") is True:
                row[3].text = "Volvió solo"
            elif ep.get("selfRecovery") is False:
                row[3].text = "Intervención"
            else:
                row[3].text = ""

    _footer(doc)
    return doc


def facility_doc(payload: dict) -> Document:
    doc = Document()
    _heading(doc, "Briefing de dirección", 0)
    _para(
        doc,
        f"Ventana {payload.get('from')} → {payload.get('to')} · "
        f"{payload.get('residentCount', 0)} residentes · "
        f"{payload.get('baselineForming', 0)} en formación de línea base · "
        f"generado {payload.get('generatedAt', '')}",
    )

    _heading(doc, "A revisar", 1)
    review = payload.get("toReview") or []
    if not review:
        _para(doc, "Nadie con tendencia negativa ni decisión pendiente en esta ventana.")
    for item in review:
        who = item.get("residentName") or item.get("residentId")
        mark = "Decisión · " if item.get("awaitingDecision") else ""
        doc.add_paragraph(f"{who} — {mark}{item.get('headline', '')}", style="List Bullet")

    _heading(doc, "Tendencias positivas", 1)
    positive = payload.get("positive") or []
    if not positive:
        _para(doc, "Sin tendencias positivas destacadas.")
    for item in positive:
        who = item.get("residentName") or item.get("residentId")
        doc.add_paragraph(f"{who} — {item.get('headline', '')}", style="List Bullet")

    residents = payload.get("residents") or []
    if residents:
        _heading(doc, "Detalle por residente", 1)
        for report in residents:
            name = report.get("residentName") or report.get("residentId")
            _heading(doc, name, 2)
            if report.get("narrative"):
                _para(doc, report["narrative"])
            for finding in report.get("findings") or []:
                if finding.get("kind") in ("POLICY", "TREND", "CLUSTER") or finding.get("polarity") == "CONCERN":
                    _para(doc, f"{finding.get('headline')}: {finding.get('body', '')}")

    _footer(doc)
    return doc


def stamp() -> str:
    return datetime.now(timezone.utc).strftime("%Y%m%d")


def save(doc: Document, name: str) -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUT_DIR / name
    doc.save(path)
    return path


def main() -> None:
    parser = argparse.ArgumentParser(description="Informes Word desde insights JSON")
    parser.add_argument("--insights-url", default=DEFAULT_INSIGHTS)
    parser.add_argument("--resident", help="id de residente (informe de ficha)")
    parser.add_argument("--facility", action="store_true", help="briefing de dirección")
    parser.add_argument("--days", type=int, default=None)
    parser.add_argument("--from", dest="date_from")
    parser.add_argument("--to", dest="date_to")
    args = parser.parse_args()
    if not args.resident and not args.facility:
        parser.error("indique --resident ID o --facility")

    q = []
    if args.days is not None:
        q.append(f"days={args.days}")
    if args.date_from:
        q.append(f"from={args.date_from}")
    if args.date_to:
        q.append(f"to={args.date_to}")
    query = f"?{'&'.join(q)}" if q else ""

    if args.resident:
        days = args.days or 30
        extra = query if q else f"?days={days}"
        payload = get_json(args.insights_url, f"/api/v1/insights/resident-chart/{args.resident}/report{extra}")
        path = save(resident_doc(payload), f"residente-{args.resident}-{stamp()}.docx")
        print(path)
        return

    days = args.days or 14
    extra = query if q else f"?days={days}"
    payload = get_json(args.insights_url, f"/api/v1/insights/facility/report{extra}")
    path = save(facility_doc(payload), f"direccion-{stamp()}.docx")
    print(path)


if __name__ == "__main__":
    main()
